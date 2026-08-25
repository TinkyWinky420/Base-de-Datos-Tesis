from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = r"C:\Users\Angel\OneDrive\Escritorio\Funciones por Rol - HiStesis.docx"

doc = Document()

# Estilo base
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Titulo
title = doc.add_heading('Funciones del Sistema por Rol - HiStesis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Documento de apoyo para estudio y presentacion de la tesis.')
doc.add_paragraph('')

# ============================================================
# ESTUDIANTE
# ============================================================
doc.add_heading('1. Rol: Estudiante', level=1)
doc.add_paragraph('El estudiante es el usuario principal del sistema. Su funcion es registrar su tesis y darle seguimiento.')
doc.add_paragraph('')

table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Funcion'
hdr[1].text = 'Que hace'

data = [
    ('Inicio', 'Pagina principal del sistema. Muestra un panel con acceso rapido a las funciones disponibles para el estudiante.'),
    ('Registrar Tesis', 'Permite crear un nuevo registro de tesis. El estudiante ingresa los datos de su tesis (numero de control, nombre del alumno, nombre de la tesis, asesor, etc.). Solo puede registrar UNA tesis por sesion.'),
    ('Revisar Mi Tesis', 'Muestra el estado de la tesis que el estudiante registro en la sesion actual. Puede ver si esta pendiente, en revision o aprobada.'),
]

for func, desc in data:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
p.add_run('El estudiante NO puede buscar tesis de otros compañeros ni ver el listado general. Solo tiene acceso a su propia tesis.')

doc.add_paragraph('')

# ============================================================
# ASESOR
# ============================================================
doc.add_heading('2. Rol: Asesor', level=1)
doc.add_paragraph('El asesor es un usuario con permisos de consulta. Puede revisar las tesis registradas por los estudiantes pero NO puede crear ni editar registros.')
doc.add_paragraph('')

table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Funcion'
hdr[1].text = 'Que hace'

data = [
    ('Buscar Tesis', 'Permite buscar tesis por diferentes criterios: numero de control, nombre del integrante, nombre del asesor o matricula. Incluye un campo de busqueda con sugerencias automaticas.'),
    ('Listado de Tesis', 'Muestra una tabla con todas las tesis registradas en el sistema. Permite ordenar y filtrar los resultados.'),
    ('Historial', 'Muestra el historial de cambios realizados en las tesis (creacion, modificaciones, eliminaciones).'),
    ('Estadisticas', 'Presenta graficas y datos estadisticos sobre las tesis registradas: por zona, por plantel, por carrera, etc.'),
]

for func, desc in data:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
p.add_run('El asesor NO puede crear, editar ni eliminar tesis. Su rol es exclusivamente de consulta y revision.')

doc.add_paragraph('')

# ============================================================
# ADMINISTRATIVO
# ============================================================
doc.add_heading('3. Rol: Administrativo', level=1)
doc.add_paragraph('El administrativo tiene los permisos mas altos del sistema. Ademas de consultar, puede gestionar la estructura organizativa (zonas, planteles y carreras).')
doc.add_paragraph('')

table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Funcion'
hdr[1].text = 'Que hace'

data = [
    ('Panel Administrativo', 'Dashboard principal que muestra un resumen general del sistema: total de tesis, zonas, planteles y carreras registradas.'),
    ('Listado de Tesis', 'Igual que el asesor: muestra la tabla completa de tesis registradas.'),
    ('Buscar Tesis', 'Igual que el asesor: busqueda por numero de control, integrante, asesor o matricula.'),
    ('Historial', 'Igual que el asesor: registro de cambios realizados en el sistema.'),
    ('Estadisticas', 'Igual que el asesor: graficas y reportes estadisticos.'),
    ('Zonas y Planteles', 'Permite crear, editar y eliminar zonas geograficas y sus planteles asociados. Es la estructura base del sistema.'),
    ('Carreras por Plantel', 'Permite gestionar las carreras disponibles en cada plantel. Asocia carreras a planteles especificos.'),
]

for func, desc in data:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = desc

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
p.add_run('El administrativo tampoco puede crear ni editar tesis. Su diferencia con el asesor es la gestion de zonas, planteles y carreras.')

doc.add_paragraph('')

# ============================================================
# RESUMEN COMPARATIVO
# ============================================================
doc.add_heading('4. Resumen Comparativo', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Funcion'
hdr[1].text = 'Estudiante'
hdr[2].text = 'Asesor'
hdr[3].text = 'Administrativo'

comparativa = [
    ('Registrar Tesis', 'Si', 'No', 'No'),
    ('Revisar Mi Tesis', 'Si', 'No', 'No'),
    ('Buscar Tesis', 'No', 'Si', 'Si'),
    ('Listado de Tesis', 'No', 'Si', 'Si'),
    ('Historial', 'No', 'Si', 'Si'),
    ('Estadisticas', 'No', 'Si', 'Si'),
    ('Zonas y Planteles', 'No', 'No', 'Si'),
    ('Carreras por Plantel', 'No', 'No', 'Si'),
]

for func, e, a, ad in comparativa:
    row = table.add_row().cells
    row[0].text = func
    row[1].text = e
    row[2].text = a
    row[3].text = ad

doc.add_paragraph('')

# ============================================================
# FLUJO DE USO
# ============================================================
doc.add_heading('5. Flujo de Uso General', level=1)

doc.add_paragraph('1. El usuario ingresa a la pagina de inicio del sistema.')
doc.add_paragraph('2. Hace clic en "Acceder" y se dirige a la pantalla de Inicio de Sesion.')
doc.add_paragraph('3. Ingresa su usuario y contrasena (los tres roles usan la contrasena: 1234).')
doc.add_paragraph('4. Segun su rol, el sistema lo redirige a la pantalla correspondiente:')
doc.add_paragraph('    - Estudiante -> Inicio', style='List Bullet')
doc.add_paragraph('    - Asesor -> Buscar Tesis', style='List Bullet')
doc.add_paragraph('    - Administrativo -> Panel Administrativo', style='List Bullet')
doc.add_paragraph('5. El usuario navega por las funciones disponibles en el menu lateral.')
doc.add_paragraph('6. Para cerrar sesion, hace clic en "Cerrar sesion" en la barra superior.')

doc.add_paragraph('')
doc.add_paragraph('')

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento de apoyo - Tesis HiStesis')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUTPUT)
print(f"Guardado: {OUTPUT}")
