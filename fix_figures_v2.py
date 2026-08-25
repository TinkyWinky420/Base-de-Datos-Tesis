import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"

shutil.copy2(HISTESIS, BACKUP)
print(f"Backup: {BACKUP}")

doc = docx.Document(HISTESIS)
total = len(doc.paragraphs)
print(f"Parrafos iniciales: {total}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def get_index_range():
    idx_start = None
    idx_end = None
    for i, p in enumerate(doc.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

# === STEP 1: Collect all body figure captions ===
print("\n[1] Analizando figuras del cuerpo...")
idx_start, idx_end = get_index_range()
print(f"  Indice: paras {idx_start}-{idx_end}")

body_figures = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        num = int(m.group(1))
        title = m.group(2)
        prev_img = has_image(doc.paragraphs[i-1]) if i > 0 else False
        cur_img = has_image(p)
        body_figures.append({
            'para_idx': i,
            'num': num,
            'title': title,
            'has_image': cur_img,
            'prev_has_image': prev_img,
            'is_duplicate': False
        })

# Show all
print(f"\n  Total figuras en cuerpo: {len(body_figures)}")
for f in body_figures:
    img_status = 'IMG' if f['has_image'] else ('PREV_IMG' if f['prev_has_image'] else 'NO_IMG')
    print(f"  [{f['para_idx']:3d}] Figura {f['num']:2d} [{img_status}] {f['title'][:55]}")

# Identify duplicates: same number, keep the one with image
seen_nums = {}
for f in body_figures:
    n = f['num']
    if n in seen_nums:
        existing = seen_nums[n]
        # If current has image and existing doesn't, mark existing as dupe
        if f['has_image'] and not existing['has_image']:
            existing['is_duplicate'] = True
            seen_nums[n] = f
        elif existing['has_image'] and not f['has_image']:
            f['is_duplicate'] = True
        else:
            # Both have or both don't have - mark second as dupe
            f['is_duplicate'] = True
    else:
        seen_nums[n] = f

dupes = [f for f in body_figures if f['is_duplicate']]
print(f"\n  Duplicados a eliminar: {len(dupes)}")
for f in dupes:
    print(f"  [{f['para_idx']}] Figura {f['num']} - {f['title'][:50]}")
