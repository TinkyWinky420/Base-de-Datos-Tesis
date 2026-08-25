import docx
from docx.oxml.ns import qn
import re

path = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
doc = docx.Document(path)

# Find the index section first
idx_start = None
idx_end = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
    if idx_start and i > idx_start + 2:
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
            idx_end = i

print(f"Indice: paras {idx_start} a {idx_end}\n")

# Now look at body figures (after the index)
print("=== FIGURAS EN EL CUERPO (despues del indice) ===\n")
body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        num = int(m.group(1))
        title = m.group(2)
        # Check if previous para has image
        prev_has_img = False
        if i > 0:
            prev = doc.paragraphs[i-1]
            for run in prev.runs:
                if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                    prev_has_img = True
                    break
        # Check current para for image
        has_img = False
        for run in p.runs:
            if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                has_img = True
                break
        
        status = "IMG" if has_img else ("IMG_ARR" if prev_has_img else "---")
        body_figs.append((i, num, title, status))
        print(f"  [{i:3d}] Figura {num:2d} [{status:7s}] {title[:60]}")

print(f"\nTotal figuras en cuerpo: {len(body_figs)}")

# Check for duplicates
nums = [f[1] for f in body_figs]
dupes = set([n for n in nums if nums.count(n) > 1])
if dupes:
    print(f"Numeros duplicados: {sorted(dupes)}")
