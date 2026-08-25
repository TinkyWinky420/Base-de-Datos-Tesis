import docx, re

def verify(path, name):
    doc = docx.Document(path)
    print(f"\n{'='*50}")
    print(f"VERIFICACION: {name}")
    print(f"{'='*50}")
    
    # Find index
    in_index = False
    print("\n--- INDICE DE FIGURAS ---")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if 'ndice de Figuras' in t:
            in_index = True
            continue
        if in_index:
            if re.match(r'^Figura\s+\d+\.', t):
                has_img = any(run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or
                              run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                              for run in p.runs)
                print(f"  {t[:80]}")
            elif t and len(t) > 3:
                if any(k in t for k in ['Cap', 'Referencia', 'Anexo']):
                    print(f"  --- FIN ---")
                    break
    
    # Find body figures
    print("\n--- CUERPO (figuras con imagen) ---")
    count = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.\s+', t):
            has_img = any(run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or
                          run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                          for run in p.runs)
            if has_img:
                count += 1
                print(f"  CON IMAGEN: {t[:80]}")
    print(f"\n  Total figuras con imagen en cuerpo: {count}")

verify(r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx", "Histesis")
verify(r"C:\Users\Angel\OneDrive\Escritorio\Manual-Tecnico.docx", "Manual Tecnico")
