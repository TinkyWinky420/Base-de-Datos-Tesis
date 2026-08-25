import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
print(f"Restaurado desde backup")

doc = docx.Document(HISTESIS)
total = len(doc.paragraphs)
print(f"Parrafos iniciales: {total}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_index():
    idx_start = idx_end = None
    for i, p in enumerate(doc.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

idx_start, idx_end = find_index()
print(f"Indice: paras {idx_start}-{idx_end}")

# === STEP 1: Collect ALL body figure paragraphs ===
body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        num = int(m.group(1))
        title = m.group(2)
        prev_img = has_image(doc.paragraphs[i-1]) if i > 0 else False
        cur_img = has_image(p)
        body_figs.append((i, num, title, cur_img, prev_img))

print(f"\n  Total figuras: {len(body_figs)}")
for idx, num, title, cur_img, prev_img in body_figs:
    s = 'IMG' if cur_img else ('PREV_IMG' if prev_img else 'NO_IMG')
    print(f"  [{idx:3d}] Figura {num:2d} [{s}] {title[:55]}")

# === STEP 2: Identify true duplicates ===
# The ONLY duplicate is Figura 12: para 717 (Relacional Logico, PREV_IMG) vs para 814 (Seguimiento Cap4, IMG)
# Also, Figura 9 appears at para 707 (no image) and 709 (has image) - same title
# Actually from the data: there's only ONE duplicate in current doc: Figura 12 at 717 and 811

# Check for actual duplicates (same number, both in body)
from collections import Counter
num_counts = Counter(f[1] for f in body_figs)
dupes = {n: c for n, c in num_counts.items() if c > 1}
print(f"\n  Numeros duplicados: {dupes}")

# === STEP 3: Delete duplicate paragraphs ===
# For each duplicated number, keep the one with image, delete the other
to_delete = []
for num, count in dupes.items():
    candidates = [(idx, title, cur_img, prev_img) for idx, n, title, cur_img, prev_img in body_figs if n == num]
    # Keep the one with image
    has_img_candidates = [(idx, t, ci, pi) for idx, t, ci, pi in candidates if ci or pi]
    no_img_candidates = [(idx, t, ci, pi) for idx, t, ci, pi in candidates if not ci and not pi]
    
    if has_img_candidates and no_img_candidates:
        # Delete the no-image ones
        for idx, t, ci, pi in no_img_candidates:
            to_delete.append(idx)
    elif len(has_img_candidates) > 1:
        # Multiple with images - keep first, delete rest
        for idx, t, ci, pi in has_img_candidates[1:]:
            to_delete.append(idx)
    elif len(no_img_candidates) > 1:
        # Multiple without images - keep first, delete rest
        for idx, t, ci, pi in no_img_candidates[1:]:
            to_delete.append(idx)

print(f"  Paras a eliminar: {to_delete}")

# Delete from highest index to lowest
for para_idx in sorted(to_delete, reverse=True):
    p = doc.paragraphs[para_idx]
    # Also delete surrounding "Nota" and image paragraphs if they're part of the figure block
    parent = p._element.getparent()
    parent.remove(p._element)
    print(f"  Borrado para {para_idx}")

# Also need to delete the image paragraph above if it's now orphaned
# and the "Nota" paragraph below

doc.save(HISTESIS)
print(f"\nDespues de borrar duplicados: saved")

# Reload and verify
doc = docx.Document(HISTESIS)
idx_start, idx_end = find_index()

# === STEP 4: Add 2 new figures ===
print("\n[2] Agregando nuevas imagenes...")

# Find "Figura 17. Diccionario de datos" to insert after its Nota
target_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Figura 17.' in t and 'Diccionario' in t:
        target_idx = i
        break

if target_idx is None:
    print("  ERROR: No se encontro Figura 17")
else:
    # Find Nota after it
    nota_idx = None
    for j in range(target_idx + 1, min(target_idx + 5, len(doc.paragraphs))):
        if 'Nota. Elaboraci' in doc.paragraphs[j].text:
            nota_idx = j
            break
    
    insert_after_idx = nota_idx if nota_idx else target_idx
    ref = doc.paragraphs[insert_after_idx]._element
    
    def make_image_elem(image_path):
        rId, _ = doc.part.get_or_add_image(image_path)
        drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(image_path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
        xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
        return etree.fromstring(xml.encode('utf-8'))
    
    def make_caption(text):
        xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
        return etree.fromstring(xml.encode('utf-8'))
    
    def make_nota():
        xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
        return etree.fromstring(xml.encode('utf-8'))
    
    # Entidad Relacion 1
    img1 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
    ref.addnext(img1)
    cap1 = make_caption("Figura 10. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
    img1.addnext(cap1)
    nota1 = make_nota()
    cap1.addnext(nota1)
    print("  + Entidad Relacion.jpeg")
    
    # Entidad Relacion 2
    img2 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
    nota1.addnext(img2)
    cap2 = make_caption("Figura 11. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
    img2.addnext(cap2)
    nota2 = make_nota()
    cap2.addnext(nota2)
    print("  + Entidad Relacion 2.jpeg")

doc.save(HISTESIS)
print("\nGuardado con nuevas imagenes")

# === STEP 5: Renumber all body figures ===
doc = docx.Document(HISTESIS)
idx_start, idx_end = find_index()

body_figs2 = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        body_figs2.append((i, int(m.group(1)), m.group(2)))

print(f"\n[3] Renumerando {len(body_figs2)} figuras...")

# Build renumber map
old_to_new = {}
for new_num, (idx, old_num, title) in enumerate(body_figs2, start=1):
    if old_num != new_num:
        old_to_new[old_num] = new_num

# Apply from highest old number to lowest
for old_num in sorted(old_to_new.keys(), reverse=True):
    new_num = old_to_new[old_num]
    # Find and replace in all paragraphs
    for p in doc.paragraphs:
        t = p.text
        if re.match(rf'^Figura\s+{old_num}\.\s', t.strip()):
            for run in p.runs:
                if f'Figura {old_num}.' in run.text:
                    run.text = run.text.replace(f'Figura {old_num}.', f'Figura {new_num}.', 1)
                    break
            break

doc.save(HISTESIS)
print(f"  Renumeracion aplicada ({len(old_to_new)} cambios)")

# === STEP 6: Update index ===
doc = docx.Document(HISTESIS)
idx_start, idx_end = find_index()

# Delete old index entries
print(f"\n[4] Actualizando indice...")
count = 0
for i in range(idx_end, idx_start, -1):
    p = doc.paragraphs[i]
    if re.match(r'^Figura\s+\d+\.', p.text.strip()):
        parent = p._element.getparent()
        parent.remove(p._element)
        count += 1
print(f"  Eliminadas {count} entradas viejas")

# Rebuild index from body figures
doc2 = docx.Document(HISTESIS)
idx_start2 = None
for i, p in enumerate(doc2.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start2 = i
        break

body_figs3 = []
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > idx_start2 + 5:  # After index area
        title = m.group(2).split('\t')[0].strip()
        body_figs3.append((int(m.group(1)), title))

ref = doc2.paragraphs[idx_start2]._element
for num, title in reversed(body_figs3):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    new_p = etree.fromstring(xml.encode('utf-8'))
    ref.addnext(new_p)

print(f"  Index reconstruido con {len(body_figs3)} entradas")

doc2.save(HISTESIS)
print(f"\n[DONE] Documento final guardado")

# Final verification
doc3 = docx.Document(HISTESIS)
print(f"\n=== VERIFICACION FINAL ===")
print(f"Parrafos totales: {len(doc3.paragraphs)}")
for i, p in enumerate(doc3.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > 100:
        pi = has_image(doc3.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        s = 'IMG' if ci else ('PREV_IMG' if pi else 'NO_IMG')
        print(f"  [{i:3d}] {t[:65]} [{s}]")
