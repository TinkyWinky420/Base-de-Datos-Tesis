import docx
from docx.oxml.ns import qn
import re

path = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
doc = docx.Document(path)

# Find index boundaries
idx_start = None
idx_end = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
    if idx_start and i > idx_start + 2:
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
            idx_end = i

print(f"Indice: paras {idx_start} a {idx_end}\n")

# Show body figures with context (2 paras before and after)
print("=== CUERPO CON CONTEXTO ===\n")
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        num = int(m.group(1))
        title = m.group(2)
        # Check images
        has_img = False
        for run in p.runs:
            if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                has_img = True
                break
        prev_has_img = False
        if i > 0:
            for run in doc.paragraphs[i-1].runs:
                if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                    prev_has_img = True
                    break
        
        status = "IMG" if has_img else ("IMG_ARR" if prev_has_img else "---")
        
        # Show context
        print(f"[{i:3d}] Figura {num:2d} [{status}] {title[:50]}")
        if i > 0:
            prev_t = doc.paragraphs[i-1].text.strip()[:60]
            print(f"      ANTES: {prev_t}")
        if i < len(doc.paragraphs)-1:
            next_t = doc.paragraphs[i+1].text.strip()[:60]
            print(f"      DESPUES: {next_t}")
        print()
