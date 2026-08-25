import docx, re

# Analyze Manual-Tecnico index structure in detail
path = r"C:\Users\Angel\OneDrive\Escritorio\Manual-Tecnico.docx"
doc = docx.Document(path)

print("=== PARAGRAPHS WITH 'Figura' ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if 'Figura' in t:
        has_tab = '\t' in t
        has_drawing = any(run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or
                          run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                          for run in p.runs)
        # Check if it's a TOC field
        has_toc = 'TOC' in str(p._element.xml)[:200]
        print(f"  [{i:3d}] tab={has_tab} img={has_drawing} toc={has_toc} : {t[:80]}")

print("\n=== INDICE SECTIONS ===")
for i, p in enumerate(doc.paragraphs):
    if 'ndice' in p.text or 'Lista' in p.text:
        print(f"  [{i:3d}] {p.text[:80]}")

print("\n=== FIGURE CAPTIONS IN BODY (with images) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        has_img = any(run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or
                      run._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                      for run in p.runs)
        if has_img:
            print(f"  [{i:3d}] {t[:80]}")
