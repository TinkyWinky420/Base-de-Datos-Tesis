import docx
from docx.oxml.ns import qn
from lxml import etree
import re

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"

import shutil
shutil.copy2(BACKUP, HISTESIS)

doc = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc.paragraphs)}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

# Find index range
idx_start = idx_end = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
    if idx_start and i > idx_start + 2:
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.', t):
            idx_end = i

print(f"Indice: paras {idx_start}-{idx_end}")

# Collect ALL paragraphs that match "Figura X." between idx_start and idx_end
index_paras = []
for i in range(idx_start + 1, idx_end + 1):
    t = doc.paragraphs[i].text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        index_paras.append(i)

print(f"Entradas indice encontradas: {len(index_paras)}")

# Check if there are MORE entries after idx_end that should also be deleted
# (old entries that weren't caught by the initial scan)
extra_end = idx_end
for i in range(idx_end + 1, min(idx_end + 60, len(doc.paragraphs))):
    t = doc.paragraphs[i].text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
        extra_end = i
        index_paras.append(i)
    elif t and not re.match(r'^Figura\s+\d+\.', t):
        # Non-figure text found - check if it's a chapter heading or similar
        if any(k in t for k in ['ndice de Tablas', 'Cap', 'Referencia', 'Anexo', 'Bibliograf', '3.', '4.', '5.']):
            break
        # Some other text - could still be old index entries below
        continue

print(f"Entradas a eliminar: {len(index_paras)}")
print(f"Rango: {index_paras[0]} a {index_paras[-1]}")

# Delete from highest to lowest
for i in sorted(index_paras, reverse=True):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

doc.save(HISTESIS)
print(f"Guardado ({len(doc.paragraphs)} parrafos)")

# Reload and rebuild index
doc = docx.Document(HISTESIS)

# Find index again
idx_start = idx_end = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
    if idx_start and i > idx_start + 2:
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.', t):
            idx_end = i

# If there are still entries, delete them
if idx_end and idx_end > idx_start:
    for i in range(idx_end, idx_start, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    doc.save(HISTESIS)
    doc = docx.Document(HISTESIS)

# Find index again
idx_start = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
        break

# Collect body figures
body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_start + 1:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > 110:
        body_figs.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

print(f"\nFiguras para indice: {len(body_figs)}")
for num, title in body_figs:
    print(f"  Figura {num}. {title[:60]}")

# Insert new index entries
ref = doc.paragraphs[idx_start]._element
for num, title in reversed(body_figs):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"\nGuardado final ({len(doc.paragraphs)} parrafos)")

# VERIFICATION
doc = docx.Document(HISTESIS)
print("\n=== INDICE ===")
in_index = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'ndice de Figuras' in t:
        in_index = True
        print(f"[{i}] {t}")
        continue
    if in_index:
        if re.match(r'^Figura\s+\d+\.', t):
            print(f"[{i}] {t[:75]}")
        elif t:
            in_index = False
            print(f"[{i}] --- FIN INDICE --- {t[:50]}")
            break

print("\n=== CUERPO ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > 110:
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        s = 'IMG' if ci else ('PREV' if pi else '---')
        print(f"[{i:3d}] {t[:70]} [{s}]")
