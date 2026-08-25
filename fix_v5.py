import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
print("Restaurado desde backup")
doc = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc.paragraphs)}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_index_range(d):
    idx_start = idx_end = None
    for i, p in enumerate(d.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

# ====================================================================
# STEP 1: Collect body figures to understand what we have
# ====================================================================
idx_s, idx_e = find_index_range(doc)
print(f"\nIndice: paras {idx_s}-{idx_e}")

body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        body_figs.append({
            'para': i,
            'old_num': int(m.group(1)),
            'title': m.group(2).split('\t')[0].strip(),
            'has_img': ci,
            'prev_has_img': pi,
        })

print(f"\nFiguras del cuerpo: {len(body_figs)}")
for f in body_figs:
    s = 'IMG' if f['has_img'] else ('PREV' if f['prev_has_img'] else '---')
    print(f"  [{f['para']:3d}] Figura {f['old_num']:2d} [{s}] {f['title'][:55]}")

# ====================================================================
# STEP 2: Delete duplicate Figura 12 (Seguimiento, para 811)
# It's a DIFFERENT figure with same number as Relacional Logico
# We keep both but renumber in step 3
# ====================================================================
# Actually, we keep ALL figures and just renumber. No deletions needed for body.
# The only "duplicate" is the number collision, not actual duplicate content.

# ====================================================================
# STEP 3: Add 2 new images after Figura 17 (Diccionario de datos)
# ====================================================================
print("\n[1] Agregando nuevas imagenes...")

# Find Figura 17 in the body
target_idx = None
for f in body_figs:
    if f['old_num'] == 17 and 'Diccionario' in f['title']:
        target_idx = f['para']
        break

# Find the Nota after it
nota_idx = None
for j in range(target_idx + 1, min(target_idx + 5, len(doc.paragraphs))):
    if 'Nota.' in doc.paragraphs[j].text:
        nota_idx = j
        break

print(f"  Figura 17 en para {target_idx}, Nota en para {nota_idx}")
ref = doc.paragraphs[nota_idx]._element

def make_image_elem(image_path):
    rId, _ = doc.part.get_or_add_image(image_path)
    drawing_xml = f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(image_path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_caption(text):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_nota():
    xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

# Insert placeholder captions (will be renumbered in step 4)
img1 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
cap1 = make_caption("Figura 99. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
nota1 = make_nota()
img2 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
cap2 = make_caption("Figura 99. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
nota2 = make_nota()

ref.addnext(img1)
img1.addnext(cap1)
cap1.addnext(nota1)
nota1.addnext(img2)
img2.addnext(cap2)
cap2.addnext(nota2)

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"  + 2 imagenes ({len(doc.paragraphs)} parrafos)")

# ====================================================================
# STEP 4: Renumber ALL body figures sequentially
# ====================================================================
print("\n[2] Renumerando figuras...")

idx_s, idx_e = find_index_range(doc)

body_figs2 = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        body_figs2.append(p)

print(f"  {len(body_figs2)} figuras a renumerar")

for new_num, p in enumerate(body_figs2, start=1):
    # Find the run that starts with "Figura X."
    renamed = False
    for run in p.runs:
        m = re.match(r'^(Figura\s+)\d+(\.)', run.text)
        if m:
            old_text = run.text
            run.text = f"Figura {new_num}." + old_text[m.end():]
            renamed = True
            break
    
    if not renamed:
        # Maybe the text is split across runs - try joining
        full = p.text
        new_full = re.sub(r'^Figura\s+\d+\.', f'Figura {new_num}.', full)
        if new_full != full:
            # Set first run to the full new text
            if p.runs:
                p.runs[0].text = new_full
                # Clear other runs
                for r in p.runs[1:]:
                    r.text = ''
    
    print(f"  -> Figura {new_num}")

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# ====================================================================
# STEP 5: Rebuild index
# ====================================================================
print("\n[3] Reconstruyendo indice...")

idx_s, idx_e = find_index_range(doc)

# Delete ALL old index entries
removed = 0
for i in range(idx_e, idx_s, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    removed += 1
print(f"  Eliminadas {removed} entradas viejas")

# Collect renumbered body figures for index
doc = docx.Document(HISTESIS)
idx_s, idx_e = find_index_range(doc)
body_figs3 = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        body_figs3.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

# Insert new index entries
ref = doc.paragraphs[idx_s]._element
for num, title in reversed(body_figs3):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"  {len(body_figs3)} entradas nuevas")

# ====================================================================
# VERIFICATION
# ====================================================================
print("\n=== VERIFICACION FINAL ===")
doc_final = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc_final.paragraphs)}")

idx_sf, idx_ef = find_index_range(doc_final)
print(f"\nIndice ({idx_sf}-{idx_ef}):")
for i in range(idx_sf, idx_ef + 1):
    t = doc_final.paragraphs[i].text.strip()
    if t:
        print(f"  {t[:75]}")

print(f"\nCuerpo:")
fig_count = 0
for i, p in enumerate(doc_final.paragraphs):
    if i <= idx_ef:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        pi = has_image(doc_final.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        s = 'IMG' if ci else ('PREV' if pi else '---')
        fig_count += 1
        print(f"  [{i:3d}] {t[:70]} [{s}]")

print(f"\nTotal figuras cuerpo: {fig_count}")
print(f"Total entradas indice: {len(body_figs3)}")
print(f"Match: {'SI' if fig_count == len(body_figs3) else 'NO'}")
