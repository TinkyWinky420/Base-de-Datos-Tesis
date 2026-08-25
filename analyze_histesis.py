import docx
from docx.oxml.ns import qn
import re

path = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
doc = docx.Document(path)

print("=== ESTRUCTURA DE FIGURAS ===\n")

# Find all figure captions
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        num = int(m.group(1))
        title = m.group(2)
        # Check if paragraph has image
        has_img = False
        for run in p.runs:
            if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                has_img = True
                break
        # Check previous paragraph for image
        prev_has_img = False
        if i > 0:
            prev = doc.paragraphs[i-1]
            for run in prev.runs:
                if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                    prev_has_img = True
                    break
        
        img_status = "CON IMAGEN" if has_img else ("IMAGEN ARRIBA" if prev_has_img else "SIN IMAGEN")
        print(f"  Figura {num:2d} [{img_status}] - {title[:70]}")

print("\n=== INDICE DE FIGURAS ===\n")
in_index = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'ndice de Figuras' in t:
        in_index = True
        continue
    if in_index:
        if re.match(r'^Figura\s+\d+\.', t):
            print(f"  {t[:80]}")
        elif t and len(t) > 3:
            if any(k in t for k in ['Cap', 'Referencia', 'Anexo', 'Bibliograf']):
                print("  --- FIN ---")
                break

print(f"\n=== RESUMEN ===")
print(f"Total parrafos: {len(doc.paragraphs)}")
print(f"Total tablas: {len(doc.tables)}")
