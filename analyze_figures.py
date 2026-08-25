import docx
from docx.oxml.ns import qn
import re

def analyze_figures(doc_path, doc_name):
    doc = docx.Document(doc_path)
    print(f"\n{'='*60}")
    print(f"ANALIZANDO: {doc_name}")
    print(f"{'='*60}")
    
    # Find all paragraphs with "Figura" 
    print("\n--- FIGURAS ENCONTRADAS ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and 'Figura' in text and ('.' in text or ':' in text):
            # Check if it has an image
            has_image = False
            for run in para.runs:
                if run._element.findall(qn('w:drawing')):
                    has_image = True
                    break
            
            # Check if next paragraph has "Nota" or image
            next_text = ""
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()[:80]
            
            img_status = "CON IMAGEN" if has_image else "SIN IMAGEN"
            print(f"  Para {i}: [{img_status}] {text[:100]}")
            if next_text and 'Nota' in next_text:
                print(f"    Siguiente: {next_text}")
    
    # Find index of figures section
    print("\n--- INDICE DE FIGURAS ---")
    in_index = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'índice de figuras' in text.lower() or 'lista de figuras' in text.lower():
            in_index = True
            print(f"  ENCONTRADO en para {i}: {text}")
            continue
        if in_index:
            if text and 'Figura' in text:
                print(f"  Para {i}: {text[:100]}")
            elif text and ('Capítulo' in text or 'Capitulo' in text or 'Referencia' in text or 'Anexo' in text):
                in_index = False
                print(f"  --- FIN INDICE (termina en para {i}: {text[:50]}) ---")
            elif not text and i > 0:
                # Check if next non-empty para exists
                pass

    # Count total paragraphs and images
    total_images = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.findall(qn('w:drawing')):
                total_images += 1
    
    # Check tables for figures too
    table_images = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if 'Figura' in para.text:
                        has_img = any(run._element.findall(qn('w:drawing')) for run in para.runs)
                        print(f"  TABLA - Figura: {para.text.strip()[:100]} (img: {has_img})")
                    for run in para.runs:
                        if run._element.findall(qn('w:drawing')):
                            table_images += 1
    
    print(f"\n  Total imagenes en body: {total_images}")
    print(f"  Total imagenes en tablas: {table_images}")

# Analyze both documents
analyze_figures(r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx", "Histesis.docx")
analyze_figures(r"C:\Users\Angel\OneDrive\Escritorio\Manual-Tecnico.docx", "Manual-Tecnico.docx")
