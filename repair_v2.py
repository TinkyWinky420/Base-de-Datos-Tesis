import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

src = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_actualizado.docx"

# Try different save approaches
doc = docx.Document(src)
print(f"Original: {len(doc.paragraphs)} parrafos")

# 1. Save as XML only (no images compression)
path1 = r"C:\Users\Angel\Documents\Histesis_v2.docx"
doc.save(path1)
print(f"v2 guardado: {os.path.getsize(path1)} bytes")

# 2. Try removing all styles that might cause issues
# Reload fresh
doc2 = docx.Document(src)

# Remove all custom styles, keep only defaults
for style in doc2.styles:
    if style.type is not None and style.name not in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        try:
            style.font.size = style.font.size
        except:
            pass

path2 = r"C:\Users\Angel\Documents\Histesis_v3.docx"
doc2.save(path2)
print(f"v3 guardado: {os.path.getsize(path2)} bytes")

# 3. Try with low-level XML fix - remove any broken customXml
from lxml import etree
import zipfile, shutil, io, copy

path3 = r"C:\Users\Angel\Documents\Histesis_v4.docx"
with zipfile.ZipFile(src, 'r') as zin:
    with zipfile.ZipFile(path3, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            # Skip customXml if broken
            if 'customXml' in item.filename:
                continue
            # Fix the relationship that points to customXml
            if item.filename == 'word/_rels/document.xml.rels':
                root = etree.fromstring(data)
                ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                to_remove = []
                for rel in root.findall('r:Relationship', ns):
                    target = rel.get('Target', '')
                    if 'customXml' in target:
                        to_remove.append(rel)
                for rel in to_remove:
                    parent = rel.getparent()
                    parent.remove(rel)
                # Also remove content type for customXml
                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            if item.filename == '[Content_Types].xml':
                root = etree.fromstring(data)
                ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'
                to_remove = []
                for override in root.findall(f'{{{ns_ct}}}Override'):
                    part = override.get('PartName', '')
                    if 'customXml' in part:
                        to_remove.append(override)
                for elem in to_remove:
                    parent = elem.getparent()
                    parent.remove(elem)
                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            zout.writestr(item, data)

print(f"v4 guardado (sin customXml): {os.path.getsize(path3)} bytes")

# Verify all
for p in [path1, path2, path3]:
    try:
        d = docx.Document(p)
        print(f"  {os.path.basename(p)}: OK ({len(d.paragraphs)} parrafos)")
    except Exception as e:
        print(f"  {os.path.basename(p)}: ERROR - {e}")
