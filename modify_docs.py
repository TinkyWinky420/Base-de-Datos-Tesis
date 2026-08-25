import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
TECNICO = r"C:\Users\Angel\OneDrive\Escritorio\Manual-Tecnico.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

def backup(path):
    b = path + ".bak"
    if not os.path.exists(b):
        shutil.copy2(path, b)

def img(name):
    return os.path.join(IMG_DIR, name)

def make_image_xml(image_path, doc):
    rId, image = doc.part.get_or_add_image(image_path)
    drawing_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="5486400" cy="4114800"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/>
        <a:graphic>
          <a:graphicFrame macro="">
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="5486400" cy="4114800"/>
            </a:xfrm>
            <a:graphicData>
              <pic:pic>
                <pic:nvPicPr>
                  <pic:cNvPr id="0" name="{os.path.basename(image_path)}"/>
                  <pic:cNvPicPr/>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{rId}"/>
                  <a:stretch><a:fillRect/></a:stretch>
                </pic:blipFill>
                <pic:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="5486400" cy="4114800"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect"/>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphicFrame>
        </a:graphic>
      </wp:inline>
    </w:drawing>'''
    return drawing_xml

def add_img_block(doc, after_para, image_path, caption):
    ref = after_para._element if hasattr(after_para, '_element') else after_para
    drawing_xml = make_image_xml(image_path, doc)
    ip_xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
    ip = etree.fromstring(ip_xml.encode('utf-8'))
    ref.addnext(ip)
    cap_xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{caption}</w:t></w:r></w:p>'
    cap = etree.fromstring(cap_xml.encode('utf-8'))
    ip.addnext(cap)
    nota_xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
    nota = etree.fromstring(nota_xml.encode('utf-8'))
    cap.addnext(nota)
    return nota

def replace_image_in_para(doc, para, image_path):
    rId, image = doc.part.get_or_add_image(image_path)
    for run in para.runs:
        for drawing in run._element.findall(qn('w:drawing')):
            for blip in drawing.findall('.//' + qn('a:blip')):
                blip.set(qn('r:embed'), rId)

def find_captions(doc):
    result = []
    for i, p in enumerate(doc.paragraphs):
        m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', p.text.strip())
        if m:
            has_img = any(run._element.findall(qn('w:drawing')) for run in p.runs)
            result.append({'idx': i, 'num': int(m.group(1)), 'title': m.group(2), 'has_img': has_img, 'para': p})
    return result

def update_index(doc, renumber_map, new_entries_list, after_fig_num=None):
    """
    Update index: renumber existing + add new entries.
    The index uses TOC fields - we modify the visible text directly.
    """
    # Find index boundaries
    start = None
    for i, p in enumerate(doc.paragraphs):
        if 'ndice de Figuras' in p.text:
            start = i + 1
            break
    if start is None:
        print("  WARNING: No se encontro indice")
        return
    
    # Find end: last paragraph with tab + "Figura N." pattern
    end = start
    for i in range(start, min(start + 80, len(doc.paragraphs))):
        t = doc.paragraphs[i].text
        if re.match(r'^Figura\s+\d+\.', t.strip()) and '\t' in t:
            end = i + 1
        elif t.strip() and any(k in t for k in ['ndice de Tablas', 'Cap', 'Referencia', 'Anexo', 'Bibliograf']):
            break
    
    print(f"  Index range: paras {start}-{end-1}")
    
    # Step 1: Renumber existing entries from highest to lowest
    for old_num in sorted(renumber_map.keys(), reverse=True):
        new_num = renumber_map[old_num]
        for i in range(start, end):
            for run in doc.paragraphs[i].runs:
                if f'Figura {old_num}.' in run.text:
                    run.text = run.text.replace(f'Figura {old_num}.', f'Figura {new_num}.', 1)
                    print(f"  Index renum: {old_num} -> {new_num}")
                    break
    
    # Step 2: Find insertion point - the paragraph to insert AFTER
    ref_para = None
    if after_fig_num is not None:
        for i in range(start, end):
            if re.match(rf'^Figura\s+{after_fig_num}\.', doc.paragraphs[i].text.strip()):
                ref_para = doc.paragraphs[i]
                print(f"  Insert after: Figura {after_fig_num} (para {i})")
                break
    
    if ref_para is None:
        # Insert after last existing entry
        for i in range(start, end):
            if re.match(r'^Figura\s+\d+\.', doc.paragraphs[i].text.strip()):
                ref_para = doc.paragraphs[i]
        if ref_para:
            print(f"  Insert after last entry (para {doc.paragraphs.index(ref_para)})")
    
    if ref_para is None:
        ref_para = doc.paragraphs[start - 1]
    
    last_elem = ref_para._element
    
    # Step 3: Insert new entries IN ORDER (not reversed)
    for num, title in new_entries_list:
        xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
        new_p = etree.fromstring(xml.encode('utf-8'))
        last_elem.addnext(new_p)
        last_elem = new_p
        print(f"  + Index: Figura {num}. {title[:50]}")


# ====================================================================
# HISTESIS
# ====================================================================
def modify_histesis():
    print("\n" + "=" * 60)
    print("HISTESIS.DOCX")
    print("=" * 60)
    backup(HISTESIS)
    doc = docx.Document(HISTESIS)
    caps = find_captions(doc)
    body = {c['num']: c for c in caps if c['num'] >= 6 and c['idx'] > 100}
    
    f9 = body.get(9)
    if f9:
        print("\n[1] Figura 9: DER.jpeg")
        last = add_img_block(doc, f9['para'], img("DER.jpeg"),
            "Figura 9. Diagrama entidad-relaci\u00f3n de HiStesis")
        
        new_figs = [
            (10, "DER simbologia.jpeg", "Simbolog\u00eda del diagrama entidad-relaci\u00f3n de HiStesis"),
            (11, "Relacional Base.jpeg", "Diagrama relacional base de la base de datos de HiStesis"),
            (12, "Relacional Logico.jpeg", "Diagrama relacional l\u00f3gico de la base de datos de HiStesis"),
            (13, "Relacional Fisico.jpeg", "Diagrama relacional f\u00edsico de la base de datos de HiStesis"),
            (14, "Flujo de datos nivel 0.jpeg", "Diagrama de flujo de datos nivel 0 de HiStesis"),
            (15, "Flujo de datos nivel 1.jpeg", "Diagrama de flujo de datos nivel 1 de HiStesis"),
            (16, "Flujo de datos nivel 2.jpeg", "Diagrama de flujo de datos nivel 2 de HiStesis"),
            (17, "Diccionario.jpeg", "Diccionario de datos de HiStesis"),
        ]
        
        for num, fname, title in new_figs:
            print(f"  Figura {num}: {title[:40]}...")
            last = add_img_block(doc, last, img(fname), f"Figura {num}. {title}")
    
    print("\n[SAVE] Guardando temporalmente...")
    doc.save(HISTESIS)
    doc = docx.Document(HISTESIS)
    
    # Renumber body captions: old 10->18, ..., old 40->48
    print("\n[R] Renumerando captions del cuerpo...")
    renumber = {old: old + 8 for old in range(10, 41)}
    for old_num in sorted(renumber.keys(), reverse=True):
        new_num = renumber[old_num]
        for p in doc.paragraphs:
            if re.match(rf'^Figura\s+{old_num}\.\s+', p.text.strip()):
                for run in p.runs:
                    if f'Figura {old_num}.' in run.text:
                        run.text = run.text.replace(f'Figura {old_num}.', f'Figura {new_num}.', 1)
                        break
                break
    
    # Renumber inline references
    print("[I] Renumerando referencias inline...")
    for old_num in sorted(renumber.keys(), reverse=True):
        new_num = renumber[old_num]
        for p in doc.paragraphs:
            if re.match(r'^Figura\s+\d+\.\s+', p.text.strip()):
                continue
            for run in p.runs:
                if run.text and f'Figura {old_num}' in run.text:
                    run.text = run.text.replace(f'Figura {old_num}', f'Figura {new_num}')
    
    # Update index
    print("\n[X] Actualizando indice...")
    new_index = [
        (10, "Simbolog\u00eda del diagrama entidad-relaci\u00f3n de HiStesis"),
        (11, "Diagrama relacional base de la base de datos de HiStesis"),
        (12, "Diagrama relacional l\u00f3gico de la base de datos de HiStesis"),
        (13, "Diagrama relacional f\u00edsico de la base de datos de HiStesis"),
        (14, "Diagrama de flujo de datos nivel 0 de HiStesis"),
        (15, "Diagrama de flujo de datos nivel 1 de HiStesis"),
        (16, "Diagrama de flujo de datos nivel 2 de HiStesis"),
        (17, "Diccionario de datos de HiStesis"),
    ]
    update_index(doc, renumber, new_index, after_fig_num=9)
    
    print("\n[GUARDANDO]...")
    doc.save(HISTESIS)
    print("  OK")


# ====================================================================
# MANUAL TECNICO
# ====================================================================
def modify_tecnico():
    print("\n" + "=" * 60)
    print("MANUAL-TECNICO.DOCX")
    print("=" * 60)
    backup(TECNICO)
    doc = docx.Document(TECNICO)
    caps = find_captions(doc)
    body = {c['num']: c for c in caps if c['idx'] > 50}
    
    f1 = body.get(1)
    
    # Figura 1: DER (replace existing image)
    print("\n[1] Figura 1: DER...")
    if f1 and f1['has_img']:
        replace_image_in_para(doc, f1['para'], img("DER.jpeg"))
        print("  OK - imagen reemplazada")
    
    # Find nota after Figura 1
    insert_point = f1['para'] if f1 else doc.paragraphs[0]
    f1_idx = f1['idx'] if f1 else 0
    for i, p in enumerate(doc.paragraphs):
        if 'Nota. Elaboraci' in p.text and abs(i - f1_idx) < 5:
            insert_point = p
            break
    
    # Figura 2: DER Simbologia (new)
    print("\n[2] Figura 2: DER Simbologia...")
    last = add_img_block(doc, insert_point, img("DER simbologia.jpeg"),
        "Figura 2. Simbolog\u00eda del diagrama entidad-relaci\u00f3n de HiStesis")
    print("  OK")
    
    # Figura 3: Relacional Base (replaces old Figura 2)
    print("\n[3] Figura 3: Relacional Base...")
    last = add_img_block(doc, last, img("Relacional Base.jpeg"),
        "Figura 3. Diagrama relacional base de la base de datos de HiStesis")
    print("  OK")
    
    # Figura 4: Relacional Logico (new)
    print("\n[4] Figura 4: Relacional Logico...")
    last = add_img_block(doc, last, img("Relacional Logico.jpeg"),
        "Figura 4. Diagrama relacional l\u00f3gico de la base de datos de HiStesis")
    print("  OK")
    
    # Figura 5: Relacional Fisico (new)
    print("\n[5] Figura 5: Relacional Fisico...")
    last = add_img_block(doc, last, img("Relacional Fisico.jpeg"),
        "Figura 5. Diagrama relacional f\u00edsico de la base de datos de HiStesis")
    print("  OK")
    
    print("\n[SAVE] Guardando temporalmente...")
    doc.save(TECNICO)
    doc = docx.Document(TECNICO)
    
    # Update index - new entries go after Figura 1 (last original before renumbering)
    # Original index has Figura 1 and Figura 2. After renumbering, Figura 2 becomes Figura 3.
    # New entries: Figura 2 (Simbologia), Figura 4 (Logico), Figura 5 (Fisico)
    # Final order: 1, 2, 3, 4, 5
    print("\n[X] Actualizando indice...")
    new_index = [
        (2, "Simbolog\u00eda del diagrama entidad-relaci\u00f3n de HiStesis"),
        (3, "Diagrama relacional base de la base de datos de HiStesis"),
        (4, "Diagrama relacional l\u00f3gico de la base de datos de HiStesis"),
        (5, "Diagrama relacional f\u00edsico de la base de datos de HiStesis"),
    ]
    update_index(doc, {}, new_index, after_fig_num=1)
    
    print("\n[GUARDANDO]...")
    doc.save(TECNICO)
    print("  OK")


if __name__ == '__main__':
    print("MODIFICACION DE DOCUMENTOS")
    print("Solo: imagenes + indice de figuras\n")
    modify_histesis()
    modify_tecnico()
    print("\n" + "=" * 60)
    print("COMPLETADO - Backups: .bak")
