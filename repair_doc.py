import docx
import os, shutil

src = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_actualizado.docx"
repair = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_actualizado_reparado.docx"
backup = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_actualizado_backup.docx"

# Backup original
shutil.copy2(src, backup)
print(f"Backup: {backup}")

# Open and re-save (repairs XML issues)
doc = docx.Document(src)
print(f"Leido: {len(doc.paragraphs)} parrafos, {len(doc.tables)} tablas")

doc.save(repair)
print(f"Guardado: {repair}")

# Verify
doc2 = docx.Document(repair)
print(f"Verificado: {len(doc2.paragraphs)} parrafos, {len(doc2.tables)} tablas")

# Compare sizes
orig = os.path.getsize(src)
newf = os.path.getsize(repair)
print(f"Original: {orig} bytes")
print(f"Reparado: {newf} bytes")
print(f"Diferencia: {newf - orig} bytes")
