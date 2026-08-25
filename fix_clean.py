import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
doc = docx.Document(HISTESIS)
print(f"Inicio: {len(doc.paragraphs)} parrafos")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

# ===== STEP 1: Find where index ends and body begins =====
# The index is a table of contents area. Find "Indice de Figuras" then
# find where the first real body paragraph starts (non-figure text after index)
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

# Scan forward from header to find where index entries end
# Index entries have tab characters, body doesn't
idx_last_entry = idx_header
for i in range(idx_header + 1, min(idx_header + 60, len(doc.paragraphs))):
    t = doc.paragraphs[i].text
    if '\t' in t and re.match(r'^Figura\s+\d+\.', t.strip()):
        idx_last_entry = i
    elif t.strip() == '' or re.match(r'^\d+\.\d+', t.strip()):
        # Empty line or section heading - might be end of index
        continue
    elif re.match(r'^Figura\s+\d+\.', t.strip()) and '\t' not in t:
        # Figure caption without tab = body figure, index ended before this
        break
    else:
        break

print(f"Indice: paras {idx_header} a {idx_last_entry}")

# ===== STEP 2: Delete index entries =====
for i in range(idx_last_entry, idx_header, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"Despues borrar indice: {len(doc.paragraphs)} parrafos")

# Re-find index header
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

# ===== STEP 3: Collect body figures =====
# Body figures are AFTER a section heading like "3.x" or "4.x" and have "Figura X."
body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_header + 1:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        body_figs.append({
            'para': i,
            'old_num': int(m.group(1)),
            'title': m.group(2).split('\t')[0].strip(),
            'has_img': ci,
            'prev_img': pi,
        })

print(f"\nFiguras cuerpo: {len(body_figs)}")
for f in body_figs:
    s = 'IMG' if f['has_img'] else ('PREV' if f['prev_img'] else '---')
    print(f"  [{f['para']:3d}] {f['old_num']:2d} [{s}] {f['title'][:55]}")

# ===== STEP 4: Add 2 new images after Figura 17 (Diccionario) =====
target = None
for f in body_figs:
    if 'Diccionario' in f['title']:
        target = f
        break

if target:
    # Find Nota after the title
    nota_idx = None
    for j in range(target['para'] + 1, min(target['para'] + 5, len(doc.paragraphs))):
        if 'Nota.' in doc.paragraphs[j].text:
            nota_idx = j
            break
    
    ref = doc.paragraphs[nota_idx]._element
    
    def make_img(path):
        rId, _ = doc.part.get_or_add_image(path)
        dwg = f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(path))%99999}" name="{os.path.basename(path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
        return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{dwg}</w:r></w:p>'.encode('utf-8'))
    
    def make_cap(text):
        return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'.encode('utf-8'))
    
    def make_nota():
        return etree.fromstring('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'.encode('utf-8'))

    img1 = make_img(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
    cap1 = make_cap("Figura XX. Entidad Relacion alternativo")
    n1 = make_nota()
    img2 = make_img(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
    cap2 = make_cap("Figura XX. Entidad Relacion variante")
    n2 = make_nota()
    
    ref.addnext(img1); img1.addnext(cap1); cap1.addnext(n1)
    n1.addnext(img2); img2.addnext(cap2); cap2.addnext(n2)
    print(f"\n+ 2 imagenes insertadas")
    doc.save(HISTESIS)
    doc = docx.Document(HISTESIS)

# ===== STEP 5: Renumber body figures =====
# Re-collect after save/reload
body_figs2 = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_header + 1:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        body_figs2.append(p)

print(f"\nRenumerando {len(body_figs2)} figuras...")
for new_num, p in enumerate(body_figs2, start=1):
    for run in p.runs:
        m = re.match(r'^(Figura\s+)\d+(\.)', run.text)
        if m:
            run.text = f"Figura {new_num}." + run.text[m.end():]
            break
    else:
        # Fallback: try all runs
        full = p.text
        new_full = re.sub(r'^Figura\s+\d+\.', f'Figura {new_num}.', full)
        if new_full != full and p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# ===== STEP 6: Build index =====
print("\nConstruyendo indice...")
# Re-find header
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

# Collect body figures for index
idx_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_header + 1:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        idx_figs.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

ref = doc.paragraphs[idx_header]._element
for num, title in reversed(idx_figs):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"{len(idx_figs)} entradas creadas")

# ===== VERIFICACION =====
doc = docx.Document(HISTESIS)
print(f"\nFinal: {len(doc.paragraphs)} parrafos")

print("\nINDICE:")
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        print(f"  [{i}] {p.text.strip()}")
        for j in range(i+1, i+50):
            t = doc.paragraphs[j].text.strip()
            if re.match(r'^Figura\s+\d+\.', t):
                print(f"  [{j}] {t[:75]}")
            elif t:
                break
        break

print("\nCUERPO:")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > 60:
        ci = has_image(p)
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        s = 'IMG' if ci else ('PREV' if pi else '---')
        print(f"  [{i:3d}] {t[:70]} [{s}]")
