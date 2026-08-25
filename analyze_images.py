import docx
from docx.oxml.ns import qn
from lxml import etree
import re

path = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
doc = docx.Document(path)

print("=== IMAGENES EN EL DOCUMENTO ===\n")
img_count = 0
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        drawings = run._element.findall(qn('w:drawing'))
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for b in blips:
                rid = b.get(qn('r:embed'))
                img_count += 1
                # Find the caption (next paragraph with "Figura")
                caption = ""
                for j in range(i+1, min(i+5, len(doc.paragraphs))):
                    t = doc.paragraphs[j].text.strip()
                    if re.match(r'^Figura\s+\d+\.', t):
                        caption = t[:60]
                        break
                print(f"  IMG {img_count:2d} at para {i}: {rid} -> {caption}")

print(f"\nTotal imagenes: {img_count}")

# Also check the rels to see image filenames
import zipfile
z = zipfile.ZipFile(path)
rels_xml = z.read('word/_rels/document.xml.rels')
root = etree.fromstring(rels_xml)
ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
print("\n=== RELACIONES DE IMAGEN ===\n")
for rel in root.findall('r:Relationship', ns):
    if 'image' in rel.get('Type', ''):
        rid = rel.get('Id')
        target = rel.get('Target')
        # Check if file exists in zip
        full_path = 'word/' + target if not target.startswith('/') else target[1:]
        exists = full_path in z.namelist()
        size = z.getinfo(full_path).file_size if exists else 0
        print(f"  {rid}: {target} ({size} bytes) {'OK' if exists else 'FALTA'}")
z.close()
