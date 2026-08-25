import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
print("Restaurado desde backup")
doc = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc.paragraphs)}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_index_range(d):
    idx_start = idx_end = None
    for i, p in enumerate(d.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

# STEP 1: Delete duplicate Figura 12 at para 811
print("\n[1] Eliminando Figura 12 duplicada...")
p811 = doc.paragraphs[811]
p811._element.getparent().remove(p811._element)
doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"  OK ({len(doc.paragraphs)} parrafos)")

# STEP 2: Add 2 new images after Figura 17 Nota
print("\n[2] Agregando nuevas imagenes...")
idx_s, idx_e = find_index_range(doc)

target_idx = None
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    if 'Figura 17.' in p.text and 'Diccionario' in p.text:
        target_idx = i
        break

nota_idx = None
for j in range(target_idx + 1, min(target_idx + 5, len(doc.paragraphs))):
    if 'Nota.' in doc.paragraphs[j].text:
        nota_idx = j
        break

ref = doc.paragraphs[nota_idx]._element

def make_image_elem(image_path):
    rId, _ = doc.part.get_or_add_image(image_path)
    drawing_xml = f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(image_path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_caption(text):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_nota():
    xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

# Use placeholder numbers, will be renumbered in step 3
img1 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
cap1 = make_caption("Figura 99. Entidad Relacion alternativo")
nota1 = make_nota()
img2 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
cap2 = make_caption("Figura 99. Entidad Relacion variante")
nota2 = make_nota()

ref.addnext(img1)
img1.addnext(cap1)
cap1.addnext(nota1)
nota1.addnext(img2)
img2.addnext(cap2)
cap2.addnext(nota2)

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"  + 2 imagenes ({len(doc.paragraphs)} parrafos)")

# STEP 3: Renumber body figures ONE BY ONE
print("\n[3] Renumerando figuras...")

idx_s, idx_e = find_index_range(doc)

# Collect ALL body figure paragraphs in document order
body_fig_paras = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.\s+', t):
        body_fig_paras.append(p)

print(f"  {len(body_fig_paras)} figuras encontradas")

# Rename each one sequentially
for new_num, p in enumerate(body_fig_paras, start=1):
    old_text = p.text
    m = re.match(r'^(Figura\s+)\d+(\.\s+.+)$', old_text)
    if m:
        new_text = f"Figura {new_num}.{m.group(2)}"
        for run in p.runs:
            if 'Figura' in run.text and '.' in run.text:
                run.text = run.text.replace(old_text, new_text, 1)
                break
        else:
            p.runs[0].text = new_text
        print(f"  -> Figura {new_num}: {new_text[10:60]}")

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# STEP 4: Rebuild index
print("\n[4] Reconstruyendo indice...")
idx_s, idx_e = find_index_range(doc)

# Delete old entries
removed = 0
for i in range(idx_e, idx_s, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    removed += 1
print(f"  Eliminadas {removed} entradas viejas")

# Collect body figures for index
doc = docx.Document(HISTESIS)
idx_s, idx_e = find_index_range(doc)
body_fig_paras2 = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        body_fig_paras2.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

ref = doc.paragraphs[idx_s]._element
for num, title in reversed(body_fig_paras2):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"  {len(body_fig_paras2)} entradas nuevas")

# VERIFICATION
print("\n=== VERIFICACION ===")
doc_final = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc_final.paragraphs)}")

idx_s, idx_e = find_index_range(doc_final)
print(f"\nIndice ({idx_s}-{idx_e}):")
for i in range(idx_s, idx_e + 1):
    t = doc_final.paragraphs[i].text.strip()
    if t:
        print(f"  {t[:75]}")

print(f"\nCuerpo:")
for i, p in enumerate(doc_final.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        pi = has_image(doc_final.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        s = 'IMG' if ci else ('PREV' if pi else '---')
        print(f"  [{i:3d}] {t[:70]} [{s}]")
