import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
print("Restaurado desde backup")
doc = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc.paragraphs)}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_index():
    idx_start = idx_end = None
    for i, p in enumerate(doc.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

idx_start, idx_end = find_index()
print(f"Indice: {idx_start}-{idx_end}")

# Show INDEX entries
print("\n=== INDICE ===")
for i in range(idx_start, idx_end + 1):
    t = doc.paragraphs[i].text.strip()
    if t:
        print(f"  [{i}] {t[:70]}")

# Show BODY figures
print("\n=== CUERPO ===")
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        s = 'IMG' if ci else ('PREV_IMG' if pi else 'NO_IMG')
        # Also show surrounding context
        print(f"  [{i:3d}] Figura {m.group(1):>2s} [{s}] {m.group(2)[:55]}")
