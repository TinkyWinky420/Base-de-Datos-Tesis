import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
print("Restaurado desde backup")
doc = docx.Document(HISTESIS)
print(f"Parrafos: {len(doc.paragraphs)}")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_index_range(d):
    idx_start = idx_end = None
    for i, p in enumerate(d.paragraphs):
        if 'ndice de Figuras' in p.text:
            idx_start = i
        if idx_start and i > idx_start + 2:
            t = p.text.strip()
            if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
                idx_end = i
    return idx_start, idx_end

# STEP 1: Delete duplicate Figura 12 at para 811
print("\n[1] Eliminando Figura 12 duplicada (para 811)...")
p811 = doc.paragraphs[811]
p811._element.getparent().remove(p811._element)
print("  OK")

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"  Parrafos despues: {len(doc.paragraphs)}")

# STEP 2: Add 2 new images after Figura 17
print("\n[2] Agregando nuevas imagenes...")

# Find Figura 17 dynamically (in the BODY, not the index)
idx_s, idx_e = find_index_range(doc)
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if i <= idx_e:
        continue
    if 'Figura 17.' in p.text and 'Diccionario' in p.text:
        target_idx = i
        break

if target_idx is None:
    print("  ERROR: No se encontro Figura 17")
    exit(1)

print(f"  Figura 17 en para {target_idx}")

# Find Nota after it
nota_idx = None
for j in range(target_idx + 1, min(target_idx + 5, len(doc.paragraphs))):
    if 'Nota.' in doc.paragraphs[j].text:
        nota_idx = j
        break

if nota_idx is None:
    print(f"  ERROR: No se encontro Nota despues de para {target_idx}")
    exit(1)

print(f"  Nota en para {nota_idx}")
ref = doc.paragraphs[nota_idx]._element

def make_image_elem(image_path):
    rId, _ = doc.part.get_or_add_image(image_path)
    drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(image_path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_caption(text):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

def make_nota():
    xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
    return etree.fromstring(xml.encode('utf-8'))

# Insert in reverse order (since addnext inserts right after)
nota2 = make_nota()
cap2 = make_caption("Figura 18. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
img2 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
nota1 = make_nota()
cap1 = make_caption("Figura 17. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
img1 = make_image_elem(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))

# Insert: nota1 -> img1 -> cap1 -> nota2 -> img2 -> cap2 after ref
ref.addnext(nota1)
nota1.addnext(img1)
img1.addnext(cap1)
cap1.addnext(nota2)
nota2.addnext(img2)
img2.addnext(cap2)
print("  + Entidad Relacion.jpeg")
print("  + Entidad Relacion 2.jpeg")

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"  Parrafos despues: {len(doc.paragraphs)}")

# STEP 3: Renumber ALL body figures sequentially (1, 2, 3, ...)
print("\n[3] Renumerando figuras del cuerpo...")

idx_start, idx_end = find_index_range(doc)

body_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        body_figs.append((i, int(m.group(1)), m.group(2)))

print(f"  Figuras encontradas: {len(body_figs)}")
for idx, num, title in body_figs:
    print(f"    [{idx:3d}] Figura {num:2d}: {title[:55]}")

# Build renumber map
old_to_new = {}
for new_num, (idx, old_num, title) in enumerate(body_figs, start=1):
    if old_num != new_num:
        old_to_new[old_num] = new_num

print(f"\n  Cambios de numeracion: {len(old_to_new)}")
for old_num in sorted(old_to_new.keys()):
    print(f"    {old_num} -> {old_to_new[old_num]}")

# Apply from highest old number to lowest (to avoid collisions)
for old_num in sorted(old_to_new.keys(), reverse=True):
    new_num = old_to_new[old_num]
    for p in doc.paragraphs:
        if re.match(rf'^Figura\s+{old_num}\.\s', p.text.strip()):
            for run in p.runs:
                if f'Figura {old_num}.' in run.text:
                    run.text = run.text.replace(f'Figura {old_num}.', f'Figura {new_num}.', 1)
                    break
            break

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# STEP 4: Rebuild index
print("\n[4] Reconstruyendo indice...")

idx_start, idx_end = find_index_range(doc)

# Delete old entries
for i in range(idx_end, idx_start, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

# Collect body figures for index
body_figs2 = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > 105:
        body_figs2.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

ref = doc.paragraphs[idx_start]._element
for num, title in reversed(body_figs2):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"  {len(body_figs2)} entradas creadas")

# FINAL VERIFICATION
print("\n=== VERIFICACION ===")
doc_final = docx.Document(HISTESIS)
idx_s, idx_e = find_index_range(doc_final)
print(f"Indice ({idx_s}-{idx_e}):")
for i in range(idx_s, idx_e + 1):
    t = doc_final.paragraphs[i].text.strip()
    if t:
        print(f"  {t[:70]}")

print(f"\nCuerpo:")
for i, p in enumerate(doc_final.paragraphs):
    if i <= idx_e:
        continue
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t):
        ci = has_image(p)
        pi = has_image(doc_final.paragraphs[i-1]) if i > 0 else False
        s = 'IMG' if ci else ('PREV_IMG' if pi else 'NO_IMG')
        print(f"  [{i:3d}] {t[:65]} [{s}]")

print(f"\nTotal parrafos: {len(doc_final.paragraphs)}")
