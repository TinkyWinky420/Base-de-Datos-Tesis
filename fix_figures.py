import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"

# Backup
shutil.copy2(HISTESIS, BACKUP)
print(f"Backup: {BACKUP}")

doc = docx.Document(HISTESIS)

# === STEP 1: Find and delete duplicate figure paragraphs (no image) ===
print("\n[1] Eliminando figuras duplicadas...")

# Find index boundaries
idx_start = None
idx_end = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i
    if idx_start and i > idx_start + 2:
        t = p.text.strip()
        if re.match(r'^Figura\s+\d+\.', t) and '\t' in t:
            idx_end = i

# Find duplicate figure paragraphs (no image, after index)
dupes_to_delete = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_end:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        # Check if has image
        has_img = False
        for run in p.runs:
            if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                has_img = True
                break
        prev_has_img = False
        if i > 0:
            for run in doc.paragraphs[i-1].runs:
                if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                    prev_has_img = True
                    break
        
        if not has_img and not prev_has_img:
            # This is a text-only figure caption - check if it's a duplicate
            num = int(m.group(1))
            title = m.group(2)
            # Check if there's another figure with same num that has image
            for j, p2 in enumerate(doc.paragraphs):
                if j <= idx_end or j == i:
                    continue
                t2 = p2.text.strip()
                m2 = re.match(rf'^Figura\s+{num}\.\s+', t2)
                if m2:
                    # Check if this one has image
                    has_img2 = False
                    for run in p2.runs:
                        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                            has_img2 = True
                            break
                    prev_has_img2 = False
                    if j > 0:
                        for run in doc.paragraphs[j-1].runs:
                            if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
                                prev_has_img2 = True
                                break
                    if has_img2 or prev_has_img2:
                        # Found duplicate - mark for deletion
                        dupes_to_delete.append(i)
                        print(f"  Borrar para {i}: Figura {num} - {title[:40]}")
                        break

# Delete duplicates (from highest to lowest to preserve indices)
for i in sorted(dupes_to_delete, reverse=True):
    p = doc.paragraphs[i]
    parent = p._element.getparent()
    parent.remove(p._element)

print(f"  Eliminadas {len(dupes_to_delete)} figuras duplicadas")

# Save and reload
doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# === STEP 2: Add 2 new figures (Entidad Relacion, Entidad Relacion 2) ===
print("\n[2] Agregando figuras nuevas...")

# Find the paragraph with "Figura 17. Diccionario de datos"
target_para = None
for i, p in enumerate(doc.paragraphs):
    if 'Figura 17.' in p.text and 'Diccionario de datos' in p.text:
        target_para = p
        break

if target_para:
    # Find the "Nota. Elaboracion propia." after it
    nota_para = None
    idx = doc.paragraphs.index(target_para)
    for j in range(idx+1, min(idx+5, len(doc.paragraphs))):
        if 'Nota. Elaboraci' in doc.paragraphs[j].text:
            nota_para = doc.paragraphs[j]
            break
    
    insert_after = nota_para if nota_para else target_para
    
    # Helper to create image paragraph
    def make_image_para(image_path):
        rId, image = doc.part.get_or_add_image(image_path)
        drawing_xml = f'''
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="5486400" cy="4114800"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:docPr id="{abs(hash(image_path)) % 99999}" name="{os.path.basename(image_path)}"/>
            <a:graphic>
              <a:graphicFrame macro="">
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="5486400" cy="4114800"/>
                </a:xfrm>
                <a:graphicData>
                  <pic:pic>
                    <pic:nvPicPr>
                      <pic:cNvPr id="0" name="{os.path.basename(image_path)}"/>
                      <pic:cNvPicPr/>
                    </pic:nvPicPr>
                    <pic:blipFill>
                      <a:blip r:embed="{rId}"/>
                      <a:stretch><a:fillRect/></a:stretch>
                    </pic:blipFill>
                    <pic:spPr>
                      <a:xfrm>
                        <a:off x="0" y="0"/>
                        <a:ext cx="5486400" cy="4114800"/>
                      </a:xfrm>
                      <a:prstGeom prst="rect"/>
                    </pic:spPr>
                  </pic:pic>
                </a:graphicData>
              </a:graphicFrame>
            </a:graphic>
          </wp:inline>
        </w:drawing>'''
        ip_xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{drawing_xml}</w:r></w:p>'
        return etree.fromstring(ip_xml.encode('utf-8'))
    
    def make_caption_para(text):
        xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
        return etree.fromstring(xml.encode('utf-8'))
    
    def make_nota_para():
        xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'
        return etree.fromstring(xml.encode('utf-8'))
    
    ref = insert_after._element
    
    # Add Entidad Relacion
    img1 = make_image_para(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
    ref.addnext(img1)
    cap1 = make_caption_para("Figura 10. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
    img1.addnext(cap1)
    nota1 = make_nota_para()
    cap1.addnext(nota1)
    print("  Agregada: Entidad Relacion.jpeg -> Figura 10")
    
    # Add Entidad Relacion 2
    img2 = make_image_para(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
    nota1.addnext(img2)
    cap2 = make_caption_para("Figura 11. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
    img2.addnext(cap2)
    nota2 = make_nota_para()
    cap2.addnext(nota2)
    print("  Agregada: Entidad Relacion 2.jpeg -> Figura 11")

# Save and reload
doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# === STEP 3: Renumber all body figures sequentially ===
print("\n[3] Renumerando figuras del cuerpo...")

# Collect all body figures in order
body_figs = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > idx_end:
        body_figs.append((i, int(m.group(1)), m.group(2), p))

# Assign new numbers
renumber_map = {}
for new_num, (idx, old_num, title, para) in enumerate(body_figs, start=1):
    if old_num != new_num:
        renumber_map[old_num] = new_num

# Apply renumbering to body captions (from highest to lowest)
for old_num in sorted(renumber_map.keys(), reverse=True):
    new_num = renumber_map[old_num]
    for p in doc.paragraphs:
        t = p.text.strip()
        if re.match(rf'^Figura\s+{old_num}\.\s+', t):
            for run in p.runs:
                if f'Figura {old_num}.' in run.text:
                    run.text = run.text.replace(f'Figura {old_num}.', f'Figura {new_num}.', 1)
                    print(f"  Figura {old_num} -> {new_num}: {t[:50]}")
                    break
            break

# Save and reload
doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# === STEP 4: Update index ===
print("\n[4] Actualizando indice de figuras...")

# Find index range again (after renumbering)
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_start = i + 1
        break

idx_end = idx_start
for i in range(idx_start, min(idx_start + 80, len(doc.paragraphs))):
    t = doc.paragraphs[i].text
    if re.match(r'^Figura\s+\d+\.', t.strip()) and '\t' in t:
        idx_end = i + 1
    elif t.strip() and any(k in t for k in ['ndice de Tablas', 'Cap', 'Referencia', 'Anexo', 'Bibliograf']):
        break

# Delete old index entries
for i in range(idx_end - 1, idx_start - 1, -1):
    p = doc.paragraphs[i]
    parent = p._element.getparent()
    parent.remove(p._element)

# Build new index entries from body figures
doc2 = docx.Document(HISTESIS)
body_figs_new = []
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > 100:  # After index area
        body_figs_new.append((int(m.group(1)), m.group(2).split('\t')[0].strip()))

# Insert new index entries after the header
ref_elem = doc.paragraphs[idx_start - 1]._element
for num, title in reversed(body_figs_new):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    new_p = etree.fromstring(xml.encode('utf-8'))
    ref_elem.addnext(new_p)

print(f"  Index actualizado con {len(body_figs_new)} entradas")

# Final save
doc.save(HISTESIS)
print("\n[OK] Documento actualizado")

# Verify
doc_final = docx.Document(HISTESIS)
print(f"\nVerificacion: {len(doc_final.paragraphs)} parrafos")
for i, p in enumerate(doc_final.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > 100:
        print(f"  {t[:70]}")
