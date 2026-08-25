import docx
import os
from lxml import etree
import zipfile

path = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_actualizado.docx"
try:
    doc = docx.Document(path)
    print(f"Abierto OK - {len(doc.paragraphs)} parrafos, {len(doc.tables)} tablas")
    
    test = r"C:\Users\Angel\OneDrive\Escritorio\test_save_check.docx"
    doc.save(test)
    print("Guardado test OK")
    os.remove(test)
    
    z = zipfile.ZipFile(path)
    doc_xml = z.read('word/document.xml')
    root = etree.fromstring(doc_xml)
    
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    ole = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
    print(f"OLE objects: {len(ole)}")
    
    vba = 'word/vbaData.xml' in z.namelist()
    print(f"VBA/Macros: {'Si' if vba else 'No'}")
    
    sigs = [n for n in z.namelist() if 'signature' in n.lower() or 'sig' in n.lower()]
    print(f"Firmas digitales: {sigs if sigs else 'No'}")
    
    z.close()
except Exception as e:
    print(f"ERROR: {type(e).__name__}: {e}")
