import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
doc = docx.Document(HISTESIS)
N = len(doc.paragraphs)
print(f"Inicio: {N} parrafos")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

def find_fig(para):
    return re.match(r'^Figura\s+(\d+)\.\s+(.+)', para.text.strip())

# === DIAGNOSTIC ===
print("\n=== ESTRUCTURA ORIGINAL ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = find_fig(p)
    if m or ('ndice de Figuras' in t):
        has_tab = '\t' in p.text
        ci = has_image(p)
        tag = 'INDICE' if has_tab else ('BODY')
        if 'ndice' in t:
            tag = 'HEADER'
        print(f"  [{i:3d}] {tag:6s} Figura {m.group(1) if m else '?':>2s}. {(m.group(2)[:50] if m else t[:50])} {'[IMG]' if ci else ''}")

# === STEP 1: Build list of paragraphs to DELETE from index ===
# Index entries have "Figura X. YYY\tNNN" pattern (with tab)
index_entries_to_delete = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '\t' in t and re.match(r'^Figura\s+\d+\.', t.strip()):
        index_entries_to_delete.append(i)

print(f"\nEntradas de indice a eliminar: {len(index_entries_to_delete)}")
print(f"  Desde para {index_entries_to_delete[0]} hasta {index_entries_to_delete[-1]}")

# Delete from highest to lowest
for i in sorted(index_entries_to_delete, reverse=True):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"Despues: {len(doc.paragraphs)} parrafos")

# === STEP 2: Add 2 new images ===
print("\nAgregando imagenes...")
target_para = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Diccionario de datos' in t and 'Figura' in t:
        target_para = i
        break

if target_para:
    # Find next Nota
    nota_para = None
    for j in range(target_para + 1, min(target_para + 5, len(doc.paragraphs))):
        if 'Nota.' in doc.paragraphs[j].text:
            nota_para = j
            break
    
    if nota_para:
        ref = doc.paragraphs[nota_para]._element
        
        def make_img(path):
            rId, _ = doc.part.get_or_add_image(path)
            dwg = f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(path))%99999}" name="{os.path.basename(path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
            return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{dwg}</w:r></w:p>'.encode('utf-8'))
        
        def make_cap(text):
            return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'.encode('utf-8'))
        
        def make_nota():
            return etree.fromstring('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'.encode('utf-8'))

        n2 = make_nota()
        c2 = make_cap("Figura XX. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
        i2 = make_img(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
        n1 = make_nota()
        c1 = make_cap("Figura XX. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
        i1 = make_img(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
        
        ref.addnext(i1); i1.addnext(c1); c1.addnext(n1)
        n1.addnext(i2); i2.addnext(c2); c2.addnext(n2)
        print("  + Entidad Relacion.jpeg")
        print("  + Entidad Relacion 2.jpeg")
        doc.save(HISTESIS)
        doc = docx.Document(HISTESIS)
    else:
        print("  ERROR: No Nota found after Diccionario")

# === STEP 3: Renumber body figures ===
print("\nRenumerando...")
body_paras = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and '\t' not in p.text:
        body_paras.append(p)

print(f"  {len(body_paras)} figuras")
for new_num, p in enumerate(body_paras, 1):
    for run in p.runs:
        m = re.match(r'^(Figura\s+)\d+(\.)', run.text)
        if m:
            run.text = f"Figura {new_num}." + run.text[m.end():]
            break
    else:
        new_text = re.sub(r'^Figura\s+\d+\.', f'Figura {new_num}.', p.text)
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# === STEP 4: Build index ===
print("\nIndice...")
idx_fig = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_fig = i
        break

# Collect body figures for index
figs_for_idx = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_fig + 1:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and '\t' not in p.text:
        figs_for_idx.append((int(m.group(1)), m.group(2).strip()))

ref = doc.paragraphs[idx_fig]._element
for num, title in reversed(figs_for_idx):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"  {len(figs_for_idx)} entradas")

# === VERIFICACION ===
doc = docx.Document(HISTESIS)
print(f"\nFinal: {len(doc.paragraphs)} parrafos\n")

print("INDICE:")
in_idx = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'ndice de Figuras' in t:
        in_idx = True
        print(f"  {t}")
        continue
    if in_idx:
        if re.match(r'^Figura\s+\d+\.', t) and '\t' not in p.text:
            print(f"  {t[:75]}")
        elif t:
            in_idx = False
            break

print("\nCUERPO:")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and '\t' not in p.text:
        ci = has_image(p)
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        s = 'IMG' if ci else ('PREV' if pi else '---')
        print(f"  [{i:3d}] {t[:70]} [{s}]")
