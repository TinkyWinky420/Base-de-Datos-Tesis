import docx
from docx.oxml.ns import qn
from lxml import etree
import re, os, shutil

HISTESIS = r"C:\Users\Angel\OneDrive\Escritorio\Histesis.docx"
BACKUP = r"C:\Users\Angel\OneDrive\Escritorio\Histesis_backup.docx"
IMG_DIR = r"C:\Users\Angel\OneDrive\Escritorio\Diagramas Tesis"

shutil.copy2(BACKUP, HISTESIS)
doc = docx.Document(HISTESIS)
print(f"Inicio: {len(doc.paragraphs)} parrafos")

def has_image(para):
    for run in para.runs:
        if run._element.findall(qn('w:drawing')) or run._element.findall(qn('w:pict')):
            return True
    return False

# STEP 1: Find header
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

# Find where body starts (first paragraph with "3." section heading after the index area)
body_start = None
for i, p in enumerate(doc.paragraphs):
    if i <= idx_header + 2:
        continue
    t = p.text.strip()
    if re.match(r'^3\.\d+', t):
        body_start = i
        break

print(f"Header: {idx_header}, Body start: {body_start}")

# STEP 2: Delete everything between header and body start (old index + ghost entries)
paras_to_delete = []
for i in range(idx_header + 1, body_start):
    t = doc.paragraphs[i].text.strip()
    if t:  # Only delete non-empty paragraphs
        paras_to_delete.append(i)

print(f"Deleting {len(paras_to_delete)} paras between header and body")
for i in sorted(paras_to_delete, reverse=True):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)
print(f"After cleanup: {len(doc.paragraphs)} paras")

# STEP 3: Collect body figures (after body_start, which shifted after deletion)
body_figs = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m and i > 5:  # well past the header
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        ci = has_image(p)
        body_figs.append({
            'para': i,
            'title': m.group(2).split('\t')[0].strip(),
            'has_img': ci,
            'prev_img': pi,
        })

print(f"\nBody figures: {len(body_figs)}")
for f in body_figs:
    s = 'IMG' if f['has_img'] else ('PREV' if f['prev_img'] else '---')
    print(f"  [{f['para']:3d}] {f['title'][:55]} [{s}]")

# STEP 4: Add 2 new images after Diccionario de datos
target = None
for f in body_figs:
    if 'Diccionario' in f['title']:
        target = f
        break

if target:
    # Find next "Nota" after Diccionario
    nota_para = None
    for j in range(target['para'] + 1, min(target['para'] + 5, len(doc.paragraphs))):
        if 'Nota.' in doc.paragraphs[j].text:
            nota_para = j
            break
    
    if nota_para is None:
        # Try to find the NEXT paragraph after the Diccionario caption
        # Sometimes there's an empty line, then Nota
        for j in range(target['para'] + 1, min(target['para'] + 8, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if 'Nota' in t or t == '':
                continue
            if re.match(r'^\d+\.\d', t) or re.match(r'^Figura', t):
                nota_para = j - 1  # insert before this
                break
    
    if nota_para is None:
        nota_para = target['para'] + 1  # just after the caption
    
    ref = doc.paragraphs[nota_para]._element
    
    def make_img(path):
        rId, _ = doc.part.get_or_add_image(path)
        dwg = f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="5486400" cy="4114800"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{abs(hash(path))%99999}" name="{os.path.basename(path)}"/><a:graphic><a:graphicFrame macro=""><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:graphicData><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{os.path.basename(path)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="4114800"/></a:xfrm><a:prstGeom prst="rect"/></pic:spPr></pic:pic></a:graphicData></a:graphicFrame></a:graphic></wp:inline></w:drawing>'''
        return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr/>{dwg}</w:r></w:p>'.encode('utf-8'))
    
    def make_cap(text):
        return etree.fromstring(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:b/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'.encode('utf-8'))
    
    def make_nota():
        return etree.fromstring('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:i/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve">Nota. Elaboraci\u00f3n propia.</w:t></w:r></w:p>'.encode('utf-8'))

    n2 = make_nota()
    c2 = make_cap("Figura XX. Diagrama entidad-relaci\u00f3n de HiStesis (variante)")
    i2 = make_img(os.path.join(IMG_DIR, "Entidad Relacion 2.jpeg"))
    n1 = make_nota()
    c1 = make_cap("Figura XX. Diagrama entidad-relaci\u00f3n de HiStesis (alternativo)")
    i1 = make_img(os.path.join(IMG_DIR, "Entidad Relacion.jpeg"))
    
    ref.addnext(i1); i1.addnext(c1); c1.addnext(n1)
    n1.addnext(i2); i2.addnext(c2); c2.addnext(n2)
    print(f"\n+ 2 imagenes insertadas despues de Diccionario (para {target['para']})")
    doc.save(HISTESIS)
    doc = docx.Document(HISTESIS)

# STEP 5: Renumber body figures
print("\nRenumerando...")
# Re-find header
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

body_paras = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > idx_header + 2:
        body_paras.append(p)

print(f"  {len(body_paras)} figuras")
for new_num, p in enumerate(body_paras, 1):
    for run in p.runs:
        m = re.match(r'^(Figura\s+)\d+(\.)', run.text)
        if m:
            run.text = f"Figura {new_num}." + run.text[m.end():]
            break
    else:
        new_text = re.sub(r'^Figura\s+\d+\.', f'Figura {new_num}.', p.text)
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''

doc.save(HISTESIS)
doc = docx.Document(HISTESIS)

# STEP 6: Build index
print("\nConstruyendo indice...")
idx_header = None
for i, p in enumerate(doc.paragraphs):
    if 'ndice de Figuras' in p.text:
        idx_header = i
        break

idx_figs = []
for i, p in enumerate(doc.paragraphs):
    if i <= idx_header + 2:
        continue
    t = p.text.strip()
    m = re.match(r'^Figura\s+(\d+)\.\s+(.+)', t)
    if m:
        idx_figs.append((int(m.group(1)), m.group(2).strip()))

ref = doc.paragraphs[idx_header]._element
for num, title in reversed(idx_figs):
    xml = f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t xml:space="preserve">Figura {num}. {title}</w:t></w:r></w:p>'
    ref.addnext(etree.fromstring(xml.encode('utf-8')))

doc.save(HISTESIS)
print(f"  {len(idx_figs)} entradas")

# VERIFICATION
doc = docx.Document(HISTESIS)
print(f"\n=== FINAL ({len(doc.paragraphs)} paras) ===")

print("\nINDICE:")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'ndice de Figuras' in t:
        print(f"  [{i}] {t}")
        for j in range(i+1, i+50):
            t2 = doc.paragraphs[j].text.strip()
            if re.match(r'^Figura\s+\d+\.', t2):
                print(f"  [{j}] {t2[:75]}")
            elif t2:
                break
        break

print("\nCUERPO:")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Figura\s+\d+\.', t) and i > 100:
        ci = has_image(p)
        pi = has_image(doc.paragraphs[i-1]) if i > 0 else False
        s = 'IMG' if ci else ('PREV' if pi else '---')
        print(f"  [{i:3d}] {t[:70]} [{s}]")
